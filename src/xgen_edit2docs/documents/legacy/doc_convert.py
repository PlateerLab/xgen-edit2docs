"""Word 97 바이너리(.doc) → DOCX 변환 — piece table + CHPX/PAPX 서식.

[MS-DOC] 스펙 (바이트 레이아웃은 reference_data/poi 의 HWPF 구현과 대조):

    WordDocument 스트림   FIB: wIdent(0xA5EC)@0, flags@0x0A(bit9 → 1Table),
                          fcPlcfBteChpx@0x00FA, fcPlcfBtePapx@0x0102,
                          fcSttbfFfn@0x0112, fcClx@0x01A2 (+lcb 각 4B 뒤)
    0Table/1Table 스트림  Clx: (Prc)* 그리고 Pcdt(0x02 + lcb + PlcPcd)
    PlcPcd                (n+1)×CP UINT32 + n×PCD(2B flags, 4B fc, 2B prm)
                          fc bit30 = 8비트 압축 조각 (offset = fc/2)

서식은 **FC(파일 오프셋) 기준**으로 붙는다 — 텍스트는 CP 순서로 조각을
이어 만들므로, 각 문자의 CP→FC 매핑을 유지한 채 bin table 을 해석한다:

    PlcfBteChpx/Papx      (n+1)×FC + n×PN(4B) — PN×512 = FKP 페이지 오프셋
    CHP FKP               [crun+1 FCs][crun×1B rgb]…[crun @511]
                          chpx = page[2*rgb]: cb(1B) + grpprl
    PAP FKP               rgb 대신 13B BX(1B offset + 12B PHE)
                          papx: cb(1B, 0 이면 다음 1B×2) + [istd 2B] + grpprl

sprm: opcode 2B (ispmd 0-8, spra 13-15) + operand (spra: 0/1→1B, 2/4/5→2B,
3→4B, 7→3B, 6→가변). 적용 sprm:
    문자  0x35 bold / 0x36 italic / 0x37 strike (토글: 1·0x81 = on),
          0x3E kul(밑줄), 0x43 hps(pt×2), 0x42 ico(팔레트)/0x70 cv(RGB),
          0x4F ftcAscii/0x50 ftcFE (SttbfFfn 글꼴 표 참조)
    문단  0x03/0x61 jc(정렬), 0x16 fInTable, 0x17 fTtp(행 끝)

표: 본문 텍스트에서 셀 끝 = 0x07 종결 문단, 행 끝 = fTtp 문단.
fInTable 연속 구간을 표로 묶어 docx 표(행/셀 격자)로 재조립한다.

특수 문자: 0x0D 문단, 0x07 셀/행 마크, 0x0B 줄바꿈, 0x0C 페이지 나눔,
0x13/0x14/0x15 필드(명령부 제거·결과 유지 — CP 매핑 보존), 0x01/0x08
개체 앵커(버림). 그림(Data 스트림 Escher)·각주는 범위 밖.
"""

from __future__ import annotations

import bisect
import io
import struct
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

from . import LegacyConvertError

_FC_COMPRESSED = 0x40000000
_FC_MASK = 0x3FFFFFFF

#: Word ico 색 팔레트 (1-16; 0 = 자동)
_ICO_RGB = {
    1: "000000", 2: "0000FF", 3: "00FFFF", 4: "00FF00", 5: "FF00FF",
    6: "FF0000", 7: "FFFF00", 8: "FFFFFF", 9: "000080", 10: "008080",
    11: "008000", 12: "800080", 13: "800000", 14: "808000", 15: "808080",
    16: "C0C0C0",
}


@dataclass
class _Piece:
    cp_start: int
    cp_end: int
    fc_raw: int

    @property
    def compressed(self) -> bool:
        return bool(self.fc_raw & _FC_COMPRESSED)

    @property
    def unit(self) -> int:
        return 1 if self.compressed else 2

    def fc_at(self, cp: int) -> int:
        base = self.fc_raw & _FC_MASK
        if self.compressed:
            return base // 2 + (cp - self.cp_start)
        return base + 2 * (cp - self.cp_start)


@dataclass
class _Chp:
    """해석된 문자 서식 (직접 서식만 — 스타일 시트 상속은 범위 밖)."""
    bold: bool = False
    italic: bool = False
    strike: bool = False
    underline: bool = False
    size_pt: Optional[float] = None
    color: Optional[str] = None
    font: Optional[str] = None
    font_fe: Optional[str] = None


@dataclass
class _TapCell:
    """TC80 (20B: rgf 2B + unused 2B + BRC80×4) + SHD80 셰이딩."""
    width_tw: int = 0
    first_merged: bool = False   # rgf 0x0001 — 가로 병합 시작
    merged: bool = False         # rgf 0x0002 — 가로 병합 연속
    vert_merge: bool = False     # rgf 0x0020 — 세로 병합 연속
    vert_restart: bool = False   # rgf 0x0040 — 세로 병합 시작
    valign: int = 0              # rgf bits7-8 — 0 top / 1 center / 2 bottom
    #: side → (val, sz_eighth_pt, color) — BRC80 그대로.
    borders: Optional[Dict[str, Tuple[str, int, str]]] = None
    shd_color: Optional[str] = None


@dataclass
class _Tap:
    """행 정의 (sprmTDefTable 0xD608 + sprmTDefTableShd 0xD609)."""
    boundaries: List[int] = field(default_factory=list)  # (itcMac+1) twips
    cells: List[_TapCell] = field(default_factory=list)


@dataclass
class _Pap:
    jc: int = 0
    in_table: bool = False
    ttp: bool = False
    space_before_tw: int = 0
    space_after_tw: int = 0
    line_mult: Optional[float] = None
    tap: Optional[_Tap] = None


@dataclass
class _DocPara:
    """종결 문자까지 확정된 문단 — runs 는 [(텍스트, cp 시작)] 조각."""
    chars: List[Tuple[str, int]] = field(default_factory=list)  # (ch, cp)
    terminator: str = "\r"
    end_cp: int = 0
    pap: _Pap = field(default_factory=_Pap)


# ── piece table ────────────────────────────────────────────────


def _read_pieces(word_stream: bytes, table_stream: bytes
                 ) -> Tuple[str, List[_Piece]]:
    if len(word_stream) < 0x01AA:
        raise LegacyConvertError("doc FIB 가 너무 짧습니다")
    (w_ident,) = struct.unpack_from("<H", word_stream, 0)
    if w_ident != 0xA5EC:
        raise LegacyConvertError("doc 시그니처(wIdent)가 아닙니다")

    fc_clx, lcb_clx = struct.unpack_from("<II", word_stream, 0x01A2)
    if lcb_clx == 0 or fc_clx + lcb_clx > len(table_stream):
        raise LegacyConvertError("doc piece table(Clx) 위치가 잘못됐습니다")
    clx = table_stream[fc_clx:fc_clx + lcb_clx]

    # Clx = (clxtGrpprl Prc)* clxtPlcPcd Pcdt
    pos = 0
    while pos < len(clx) and clx[pos] == 0x01:
        (cb,) = struct.unpack_from("<H", clx, pos + 1)
        pos += 3 + cb
    if pos >= len(clx) or clx[pos] != 0x02:
        raise LegacyConvertError("doc Pcdt 마커가 없습니다")
    (lcb,) = struct.unpack_from("<I", clx, pos + 1)
    plc = clx[pos + 5:pos + 5 + lcb]
    if len(plc) < 4 or (len(plc) - 4) % 12 != 0:
        raise LegacyConvertError("doc PlcPcd 크기가 어긋납니다")
    n = (len(plc) - 4) // 12
    cps = struct.unpack_from(f"<{n + 1}I", plc, 0)
    out: List[str] = []
    pieces: List[_Piece] = []
    for i in range(n):
        _flags, fc_raw, _prm = struct.unpack_from("<HIH", plc, 4 * (n + 1) + 8 * i)
        count = cps[i + 1] - cps[i]
        if count <= 0:
            continue
        piece = _Piece(cps[i], cps[i + 1], fc_raw)
        pieces.append(piece)
        if piece.compressed:
            off = (fc_raw & _FC_MASK) // 2
            raw = word_stream[off:off + count]
            out.append(raw.decode("cp1252", errors="replace"))
        else:
            off = fc_raw & _FC_MASK
            raw = word_stream[off:off + count * 2]
            out.append(raw.decode("utf-16le", errors="replace"))
    return "".join(out), pieces


# ── FKP bin tables ─────────────────────────────────────────────


def _read_bin_table(word_stream: bytes, table_stream: bytes,
                    fc: int, lcb: int, kind: str
                    ) -> Tuple[List[int], List[bytes]]:
    """PlcfBteChpx/Papx → 정렬된 (fc 경계 목록, 경계별 grpprl).

    반환: (starts, grpprls) — starts[i] ≤ fc < starts[i+1] 구간의 서식이
    grpprls[i]. 깨진 페이지는 건너뛴다 (서식만 잃고 텍스트는 무사).
    """
    starts: List[int] = []
    grpprls: List[bytes] = []
    if lcb < 8 or fc + lcb > len(table_stream):
        return starts, grpprls
    n = (lcb - 4) // 8
    try:
        fcs = struct.unpack_from(f"<{n + 1}I", table_stream, fc)
        pns = struct.unpack_from(f"<{n}I", table_stream, fc + 4 * (n + 1))
    except struct.error:
        return starts, grpprls
    del fcs  # bin table 의 FC 경계는 FKP 안의 rgfc 가 더 정밀하다
    for pn in pns:
        page_off = (pn & 0x3FFFFF) * 512
        page = word_stream[page_off:page_off + 512]
        if len(page) < 512:
            continue
        crun = page[511]
        if crun == 0 or (crun + 1) * 4 > 511:
            continue
        try:
            rgfc = struct.unpack_from(f"<{crun + 1}I", page, 0)
        except struct.error:
            continue
        for i in range(crun):
            if kind == "chp":
                rgb = page[(crun + 1) * 4 + i]
                if rgb == 0:
                    grpprl = b""
                else:
                    off = 2 * rgb
                    if off >= 511:
                        continue
                    cb = page[off]
                    grpprl = page[off + 1:off + 1 + cb]
            else:  # pap — BX 13B, papx 는 istd 2B 로 시작
                bx_off = (crun + 1) * 4 + i * 13
                if bx_off >= 511:
                    continue
                rgb = page[bx_off]
                if rgb == 0:
                    grpprl = b""
                else:
                    off = 2 * rgb
                    if off >= 511:
                        continue
                    size = 2 * page[off]
                    if size == 0:
                        off += 1
                        size = 2 * page[off]
                    else:
                        size -= 1
                    grpprl = page[off + 1:off + 1 + size]
                    grpprl = grpprl[2:]  # istd 건너뜀
            starts.append(rgfc[i])
            grpprls.append(grpprl)
            # 구간 끝 경계도 넣어 두면 lookup 이 단순해진다 — 마지막에 정렬
    # fc 순 정렬 (FKP 페이지가 fc 순이 아닐 수 있다)
    order = sorted(range(len(starts)), key=lambda k: starts[k])
    return [starts[k] for k in order], [grpprls[k] for k in order]


def _lookup(starts: List[int], grpprls: List[bytes], fc: int) -> bytes:
    if not starts:
        return b""
    idx = bisect.bisect_right(starts, fc) - 1
    return grpprls[idx] if idx >= 0 else b""


# ── sprm 해석 ──────────────────────────────────────────────────


def _iter_sprms(grpprl: bytes):
    """(opcode, operand bytes) — [MS-DOC] 2.6.1, POI SprmOperation 대조."""
    pos, n = 0, len(grpprl)
    while pos + 2 <= n:
        (opcode,) = struct.unpack_from("<H", grpprl, pos)
        pos += 2
        spra = opcode >> 13
        if spra in (0, 1):
            size = 1
        elif spra in (2, 4, 5):
            size = 2
        elif spra == 3:
            size = 4
        elif spra == 7:
            size = 3
        else:  # 6 — 가변: 길이 바이트 (표 sprm 0xD608/0xC615 는 2B 길이)
            if opcode in (0xD608, 0xC615):
                if pos + 2 > n:
                    return
                (size,) = struct.unpack_from("<H", grpprl, pos)
                pos += 2
                size = max(0, size - 1)
            else:
                if pos + 1 > n:
                    return
                size = grpprl[pos]
                pos += 1
        if pos + size > n:
            return
        yield opcode, grpprl[pos:pos + size]
        pos += size


def _toggle_on(operand: bytes) -> Optional[bool]:
    """토글 sprm: 0 off / 1 on / 0x80 스타일따름 / 0x81 스타일반전.

    직접 서식만 보므로 0x80 은 무시(None), 0x81 은 on 취급 (기본 스타일
    off 가정 — POI getCHPFlag 와 동일한 근사)."""
    if not operand:
        return None
    v = operand[0]
    if v == 0:
        return False
    if v == 1 or v == 0x81:
        return True
    return None


def _chp_of(grpprl: bytes, fonts: List[str]) -> _Chp:
    chp = _Chp()
    for opcode, operand in _iter_sprms(grpprl):
        op = opcode & 0x1FF
        if op == 0x35:
            v = _toggle_on(operand)
            if v is not None:
                chp.bold = v
        elif op == 0x36:
            v = _toggle_on(operand)
            if v is not None:
                chp.italic = v
        elif op == 0x37:
            v = _toggle_on(operand)
            if v is not None:
                chp.strike = v
        elif op == 0x3E and operand:
            chp.underline = operand[0] != 0
        elif op == 0x43 and len(operand) >= 2:
            (hps,) = struct.unpack_from("<H", operand, 0)
            if 2 <= hps <= 3276:
                chp.size_pt = hps / 2.0
        elif op == 0x42 and operand:
            chp.color = _ICO_RGB.get(operand[0])
        elif op == 0x70 and len(operand) >= 4:
            r, g, b, auto = operand[0], operand[1], operand[2], operand[3]
            if auto != 0xFF:
                chp.color = f"{r:02X}{g:02X}{b:02X}"
        elif op == 0x4F and len(operand) >= 2:
            (ftc,) = struct.unpack_from("<H", operand, 0)
            if 0 <= ftc < len(fonts):
                chp.font = fonts[ftc]
        elif op == 0x50 and len(operand) >= 2:
            (ftc,) = struct.unpack_from("<H", operand, 0)
            if 0 <= ftc < len(fonts):
                chp.font_fe = fonts[ftc]
    return chp


#: BRC80 brcType → docx 테두리 val (0/255 = 없음)
_BRC_VAL = {0: "nil", 255: "nil", 1: "single", 2: "single", 3: "double",
            5: "single", 6: "dotted", 7: "dashed", 8: "dotDash",
            9: "dotDotDash", 10: "triple"}


def _parse_brc(operand: bytes, off: int) -> Tuple[str, int, str]:
    """BRC80 4B: w1 = dpt(1/8pt, 0xFF) | brcType<<8, w2 = ico(0xFF)."""
    w1, w2 = struct.unpack_from("<2H", operand, off)
    dpt = w1 & 0xFF
    brc_type = (w1 >> 8) & 0xFF
    ico = w2 & 0xFF
    val = _BRC_VAL.get(brc_type, "single")
    color = _ICO_RGB.get(ico, "000000")
    return val, max(2, min(dpt, 48)), color


def _parse_tdef(operand: bytes) -> Optional[_Tap]:
    """sprmTDefTable: itcMac(1B) + (itcMac+1)×INT16 경계(twips) +
    itcMac×TC80(20B — 있을 때만)."""
    if not operand:
        return None
    itc = operand[0]
    if itc == 0 or itc > 63:
        return None
    need = 1 + (itc + 1) * 2
    if len(operand) < need:
        return None
    tap = _Tap()
    tap.boundaries = list(struct.unpack_from(f"<{itc + 1}h", operand, 1))
    tc_start = need
    for i in range(itc):
        cell = _TapCell()
        if i + 1 <= itc and len(tap.boundaries) > i + 1:
            cell.width_tw = max(0, tap.boundaries[i + 1] - tap.boundaries[i])
        off = tc_start + i * 20
        if off + 20 <= len(operand):
            (rgf,) = struct.unpack_from("<H", operand, off)
            cell.first_merged = bool(rgf & 0x0001)
            cell.merged = bool(rgf & 0x0002)
            cell.vert_merge = bool(rgf & 0x0020)
            cell.vert_restart = bool(rgf & 0x0040)
            cell.valign = (rgf >> 7) & 0x3
            cell.borders = {
                "top": _parse_brc(operand, off + 4),
                "left": _parse_brc(operand, off + 8),
                "bottom": _parse_brc(operand, off + 12),
                "right": _parse_brc(operand, off + 16),
            }
        tap.cells.append(cell)
    return tap


def _pap_of(grpprl: bytes) -> _Pap:
    pap = _Pap()
    shd_raw: Optional[bytes] = None
    for opcode, operand in _iter_sprms(grpprl):
        op = opcode & 0x1FF
        if op in (0x03, 0x61) and operand:  # sprmPJc (97) / sprmPJc80 (2000+)
            pap.jc = operand[0]
        elif op == 0x16 and operand:
            pap.in_table = operand[0] != 0
        elif op == 0x17 and operand:
            pap.ttp = operand[0] != 0
        elif op == 0x12 and len(operand) >= 4:  # sprmPDyaLine — LSPD
            dya, f_mult = struct.unpack_from("<hh", operand, 0)
            if f_mult and 60 <= dya <= 1200:
                pap.line_mult = dya / 240.0
        elif op == 0x13 and len(operand) >= 2:  # sprmPDyaBefore (twips)
            (pap.space_before_tw,) = struct.unpack_from("<H", operand, 0)
        elif op == 0x14 and len(operand) >= 2:  # sprmPDyaAfter
            (pap.space_after_tw,) = struct.unpack_from("<H", operand, 0)
        elif op == 0x08 and opcode == 0xD608:  # sprmTDefTable
            tap = _parse_tdef(operand)
            if tap is not None:
                pap.tap = tap
        elif op == 0x09 and opcode == 0xD609:  # sprmTDefTableShd
            shd_raw = operand
    if pap.tap is not None and shd_raw:
        # SHD80 2B: icoFore bits0-4, icoBack bits5-9, ipat bits10-15
        for i, cell in enumerate(pap.tap.cells):
            off = i * 2
            if off + 2 > len(shd_raw):
                break
            (shd,) = struct.unpack_from("<H", shd_raw, off)
            ico_fore = shd & 0x1F
            ico_back = (shd >> 5) & 0x1F
            ipat = (shd >> 10) & 0x3F
            if ipat == 1 and ico_fore:  # solid 전경색
                cell.shd_color = _ICO_RGB.get(ico_fore)
            elif ico_back:
                cell.shd_color = _ICO_RGB.get(ico_back)
    return pap


# ── 글꼴 표 (SttbfFfn) ────────────────────────────────────────


def _read_font_names(table_stream: bytes, fc: int, lcb: int) -> List[str]:
    """FFN 목록 — 이름은 엔트리 시작 +40 의 UTF-16 널종결 ([MS-DOC] 2.9.63)."""
    names: List[str] = []
    if lcb < 6 or fc + lcb > len(table_stream):
        return names
    try:
        marker, count = struct.unpack_from("<HH", table_stream, fc)
        pos = fc + 6  # marker + count + cbExtra
        if marker != 0xFFFF:  # 확장 STTB 가 아니면 해석 포기 (이름만 잃는다)
            return names
        for _ in range(count):
            if pos >= fc + lcb:
                break
            cb_m1 = table_stream[pos]
            entry = table_stream[pos + 1:pos + 1 + cb_m1]
            name = ""
            if len(entry) > 39:
                raw = entry[39:]
                name = raw.decode("utf-16le", errors="replace").split("\x00")[0]
            names.append(name)
            pos += 1 + cb_m1
    except (struct.error, IndexError):
        pass
    return names


# ── 필드 제거 (CP 매핑 보존) ──────────────────────────────────


def _visible_chars(text: str) -> List[Tuple[str, int]]:
    """필드 명령부(0x13..0x14)를 지운 (문자, 원본 CP) 목록."""
    out: List[Tuple[str, int]] = []
    depth_cmd = 0
    for cp, ch in enumerate(text):
        code = ord(ch)
        if code == 0x13:
            depth_cmd += 1
            continue
        if code == 0x14:
            if depth_cmd > 0:
                depth_cmd -= 1
            continue
        if code == 0x15:
            continue
        if depth_cmd > 0:
            continue
        out.append((ch, cp))
    return out


# ── 본문 → 문단/표 ────────────────────────────────────────────


def _split_paragraphs(chars: List[Tuple[str, int]]) -> List[_DocPara]:
    paras: List[_DocPara] = []
    cur = _DocPara()
    for ch, cp in chars:
        if ch in ("\r", "\x07"):
            cur.terminator = ch
            cur.end_cp = cp
            paras.append(cur)
            cur = _DocPara()
        else:
            cur.chars.append((ch, cp))
    if cur.chars:
        cur.end_cp = cur.chars[-1][1]
        paras.append(cur)
    return paras


def doc_to_docx(content: bytes) -> bytes:
    import olefile
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if not olefile.isOleFile(io.BytesIO(content)):
        raise LegacyConvertError("doc 가 아닙니다 (OLE 복합문서 아님)")
    ole = olefile.OleFileIO(io.BytesIO(content))
    try:
        if not ole.exists("WordDocument"):
            raise LegacyConvertError("doc WordDocument 스트림이 없습니다")
        word_stream = ole.openstream("WordDocument").read()
        (flags,) = struct.unpack_from("<H", word_stream, 0x0A)
        table_name = "1Table" if (flags & 0x0200) else "0Table"
        if not ole.exists(table_name):
            # 저장 도중 죽은 파일 등 — 반대쪽이라도 있으면 그걸 쓴다.
            table_name = "0Table" if table_name == "1Table" else "1Table"
            if not ole.exists(table_name):
                raise LegacyConvertError("doc Table 스트림이 없습니다")
        table_stream = ole.openstream(table_name).read()
    finally:
        ole.close()

    text, pieces = _read_pieces(word_stream, table_stream)

    def fib_fclcb(off: int) -> Tuple[int, int]:
        if off + 8 <= len(word_stream):
            return struct.unpack_from("<II", word_stream, off)
        return 0, 0

    fc_chp, lcb_chp = fib_fclcb(0x00FA)
    fc_pap, lcb_pap = fib_fclcb(0x0102)
    fc_ffn, lcb_ffn = fib_fclcb(0x0112)

    chp_starts, chp_grpprls = _read_bin_table(
        word_stream, table_stream, fc_chp, lcb_chp, "chp")
    pap_starts, pap_grpprls = _read_bin_table(
        word_stream, table_stream, fc_pap, lcb_pap, "pap")
    fonts = _read_font_names(table_stream, fc_ffn, lcb_ffn)

    piece_starts = [p.cp_start for p in pieces]

    def piece_of(cp: int) -> Optional[_Piece]:
        idx = bisect.bisect_right(piece_starts, cp) - 1
        if idx < 0:
            return None
        p = pieces[idx]
        return p if p.cp_start <= cp < p.cp_end else None

    chp_cache: Dict[bytes, _Chp] = {}

    def chp_at(cp: int) -> _Chp:
        p = piece_of(cp)
        if p is None:
            return _Chp()
        grpprl = _lookup(chp_starts, chp_grpprls, p.fc_at(cp))
        chp = chp_cache.get(grpprl)
        if chp is None:
            chp = chp_cache[grpprl] = _chp_of(grpprl, fonts)
        return chp

    def pap_at(cp: int) -> _Pap:
        p = piece_of(cp)
        if p is None:
            return _Pap()
        return _pap_of(_lookup(pap_starts, pap_grpprls, p.fc_at(cp)))

    paras = _split_paragraphs(_visible_chars(text))
    for para in paras:
        para.pap = pap_at(para.end_cp)

    # ── DOCX 조립 ──
    doc = Document()
    _WD_JC = {0: WD_ALIGN_PARAGRAPH.LEFT, 1: WD_ALIGN_PARAGRAPH.CENTER,
              2: WD_ALIGN_PARAGRAPH.RIGHT, 3: WD_ALIGN_PARAGRAPH.JUSTIFY}

    def emit_para(target, para: _DocPara) -> None:
        if para.pap.jc in _WD_JC and para.pap.jc != 0:
            target.alignment = _WD_JC[para.pap.jc]
        pf = target.paragraph_format
        if para.pap.line_mult is not None and abs(para.pap.line_mult - 1.0) > 0.01:
            pf.line_spacing = para.pap.line_mult
        if para.pap.space_before_tw > 0:
            pf.space_before = Pt(min(para.pap.space_before_tw, 2880) / 20.0)
        if para.pap.space_after_tw > 0:
            pf.space_after = Pt(min(para.pap.space_after_tw, 2880) / 20.0)
        # 문자 서식 경계로 런을 쪼갠다 — 같은 _Chp 가 이어지면 한 런.
        run_chars: List[str] = []
        run_chp: Optional[_Chp] = None
        run_obj = None

        def flush() -> None:
            nonlocal run_chars
            run_chars = []

        def current_run():
            nonlocal run_obj
            if run_obj is None:
                run_obj = target.add_run()
                st = run_chp or _Chp()
                if st.size_pt:
                    run_obj.font.size = Pt(st.size_pt)
                if st.bold:
                    run_obj.bold = True
                if st.italic:
                    run_obj.italic = True
                if st.underline:
                    run_obj.underline = True
                if st.strike:
                    run_obj.font.strike = True
                if st.color:
                    run_obj.font.color.rgb = RGBColor.from_string(st.color)
                name = st.font or st.font_fe
                if name:
                    run_obj.font.name = name
                    rPr = run_obj._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn("w:eastAsia"), st.font_fe or name)
            return run_obj

        def emit_text() -> None:
            nonlocal run_obj
            if run_chars:
                current_run().add_text("".join(run_chars))
            flush()

        for ch, cp in para.chars:
            code = ord(ch)
            chp = chp_at(cp)
            if run_chp is not None and chp is not run_chp and run_chars:
                emit_text()
                run_obj = None
            run_chp = chp
            if code == 0x0B:
                emit_text()
                current_run().add_break()
            elif code == 0x0C:
                emit_text()
                current_run().add_break(WD_BREAK.PAGE)
            elif code in (0x01, 0x08, 0x00, 0x1F):
                continue  # 개체 앵커/소프트하이픈 — 버림
            elif code == 0x1E:
                run_chars.append("-")
            else:
                run_chars.append(ch)
        emit_text()

    i = 0
    n_paras = len(paras)
    while i < n_paras:
        para = paras[i]
        if not para.pap.in_table:
            p = doc.add_paragraph()
            emit_para(p, para)
            i += 1
            continue
        # ── 표 구간: fInTable 연속 문단 → 행(fTtp 경계)/셀(0x07 종결) ──
        rows: List[List[List[_DocPara]]] = []
        row_taps: List[Optional[_Tap]] = []
        cur_row: List[List[_DocPara]] = []
        cur_cell: List[_DocPara] = []
        while i < n_paras and paras[i].pap.in_table:
            tp = paras[i]
            if tp.pap.ttp:
                if cur_cell:
                    cur_row.append(cur_cell)
                    cur_cell = []
                if cur_row:
                    rows.append(cur_row)
                    row_taps.append(tp.pap.tap)  # 행 정의는 행 끝 문단에 실린다
                    cur_row = []
            elif tp.terminator == "\x07":
                cur_cell.append(tp)
                cur_row.append(cur_cell)
                cur_cell = []
            else:
                cur_cell.append(tp)
            i += 1
        if cur_cell:
            cur_row.append(cur_cell)
        if cur_row:
            rows.append(cur_row)
            row_taps.append(None)
        if not rows:
            continue
        n_cols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=n_cols)
        try:
            tbl.style = "Table Grid"
        except Exception:  # noqa: BLE001
            pass
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Emu

        # 열 너비 — 첫 행 정의의 경계(twips) 차분
        first_tap = next((t for t in row_taps if t is not None), None)
        if first_tap is not None:
            for j, column in enumerate(tbl.columns):
                if j < len(first_tap.cells) and first_tap.cells[j].width_tw > 0:
                    try:
                        column.width = Emu(first_tap.cells[j].width_tw * 635)
                    except Exception:  # noqa: BLE001
                        pass

        def set_borders(cell_obj, borders) -> None:
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            tcb = tc_pr.find(qn("w:tcBorders"))
            if tcb is None:
                tcb = OxmlElement("w:tcBorders")
                tc_pr.append(tcb)
            for name, (val, sz, color) in borders.items():
                el = tcb.find(qn(f"w:{name}"))
                if el is None:
                    el = OxmlElement(f"w:{name}")
                    tcb.append(el)
                el.set(qn("w:val"), val)
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:color"), color)

        def set_shd(cell_obj, rgb: str) -> None:
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), rgb)
            tc_pr.append(shd)

        def set_valign(cell_obj, v: int) -> None:
            if v not in (1, 2):
                return
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            va = OxmlElement("w:vAlign")
            va.set(qn("w:val"), "center" if v == 1 else "bottom")
            tc_pr.append(va)

        # TC80 rgf 기반 병합 — 가로(fFirstMerged+fMerged 연속),
        # 세로(fVertRestart 앵커 + 아래 행 fVertMerge 연속, 같은 열 색인).
        for r_i, tap in enumerate(row_taps[:len(rows)]):
            if tap is None:
                continue
            c = 0
            while c < min(len(tap.cells), n_cols):
                if tap.cells[c].first_merged:
                    end = c
                    while (end + 1 < min(len(tap.cells), n_cols)
                           and tap.cells[end + 1].merged):
                        end += 1
                    if end > c:
                        try:
                            tbl.cell(r_i, c).merge(tbl.cell(r_i, end))
                        except Exception:  # noqa: BLE001
                            pass
                    c = end + 1
                else:
                    c += 1
        for c in range(n_cols):
            r_i = 0
            while r_i < len(rows):
                tap = row_taps[r_i] if r_i < len(row_taps) else None
                if (tap is not None and c < len(tap.cells)
                        and tap.cells[c].vert_restart):
                    end = r_i
                    while end + 1 < len(rows):
                        nt = row_taps[end + 1] if end + 1 < len(row_taps) else None
                        if (nt is not None and c < len(nt.cells)
                                and nt.cells[c].vert_merge
                                and not nt.cells[c].vert_restart):
                            end += 1
                        else:
                            break
                    if end > r_i:
                        try:
                            tbl.cell(r_i, c).merge(tbl.cell(end, c))
                        except Exception:  # noqa: BLE001
                            pass
                    r_i = end + 1
                else:
                    r_i += 1

        for r_i, row in enumerate(rows):
            tap = row_taps[r_i] if r_i < len(row_taps) else None
            for c_i, cell_paras in enumerate(row[:n_cols]):
                try:
                    cell = tbl.cell(r_i, c_i)
                except Exception:  # noqa: BLE001
                    continue
                if tap is not None and c_i < len(tap.cells):
                    tc = tap.cells[c_i]
                    # 병합 연속 셀은 앵커와 같은 _tc 를 공유한다 — 여기서
                    # 스타일을 다시 쓰면 앵커의 테두리(예: 상변 이중선)를
                    # 연속 행 값으로 덮어써 버린다.
                    continuation = (tc.merged and not tc.first_merged) or (
                        tc.vert_merge and not tc.vert_restart)
                    if not continuation:
                        if tc.borders:
                            set_borders(cell, tc.borders)
                        if tc.shd_color:
                            set_shd(cell, tc.shd_color)
                        set_valign(cell, tc.valign)
                first = True
                for cp_para in cell_paras:
                    if first and not cell.paragraphs[0].runs:
                        cell.paragraphs[0].text = ""
                        target = cell.paragraphs[0]
                        first = False
                    else:
                        first = False
                        target = cell.add_paragraph()
                    emit_para(target, cp_para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
