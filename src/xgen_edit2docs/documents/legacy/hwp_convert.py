"""HWP 5.0 (한글 바이너리) → DOCX 구조 보존 변환.

HWP 5.0 은 OLE 복합문서다 (레퍼런스: 한컴 공개 스펙 白書 —
reference_data/hwpml_3.0_spec.pdf 의 구조 의미 + pyhwp binmodel 의 바이트
레이아웃, 표 번호는 한컴 스펙 기준):

    FileHeader          32B 시그니처 + version + flags(bit0 압축, bit1 암호)
    DocInfo             레코드 스트림 — ID_MAPPINGS/FACE_NAME/BORDER_FILL/
                        CHAR_SHAPE/PARA_SHAPE/BIN_DATA …
    BodyText/Section0…  레코드 스트림 — PARA_HEADER/PARA_TEXT/
                        PARA_CHAR_SHAPE/CTRL_HEADER/LIST_HEADER/TABLE/
                        SHAPE_COMPONENT(_PICTURE)/PAGE_DEF
    BinData/BIN%04X.ext 삽입 이미지 원본 (압축 플래그를 따라간다)

레코드 헤더는 UINT32 하나: tagid(10b) | level(10b) | size(12b, 0xFFF 이면
다음 UINT32 가 실제 크기). 압축 플래그가 켜진 스트림은 raw zlib(-15).
레코드의 부모-자식 관계는 level 로 표현된다 — 여기서는 트리로 복원한 뒤
의미 단위(문단/표/그림/글상자/머리말)로 해석한다.

본문 텍스트는 UTF-16LE 이되 0x00~0x1F 코드가 컨트롤 문자다 — 문자형(1워드)/
인라인·확장형(8워드) 크기 표에 따라 건너뛴다 (pyhwp ControlChar 표와 동일).

충실도 범위:
- 문단: 정렬(PARA_SHAPE align), 런 스타일(글꼴/크기/굵게/기울임/밑줄/
  취소선/색 — CHAR_SHAPE 표 28/30)
- 표(표 70~75): rows×cols 격자, **병합(colspan/rowspan → gridSpan/vMerge)**,
  열 너비/행 높이, 셀 배경·**변별 테두리(표 18/20/21 — 실선/대시/없음,
  굵기, 색)**, 셀 수직 정렬(표 60 listflags), 셀 안쪽 여백(표 75 → tcMar),
  셀 안 문단 전체 스타일, 중첩 표(셀 안 표 — 재귀), 캡션
- 문단 간격: 줄간격(표 38 RATIO)·문단 앞/뒤 간격 → w:spacing
- 그림(표 102): BinData 임베딩 → docx 인라인 이미지 (개체 요소 크기 반영)
- 글상자: 텍스트를 본문 문단으로 (위치는 범위 밖 — 내용 유실 없음)
- 머리말/꼬리말: 첫 정의를 docx 섹션 header/footer 텍스트로
- 페이지: PAGE_DEF 크기/여백

수식·도형 좌표·각주는 범위 밖. 암호/배포용 문서는 명시 거부.
"""

from __future__ import annotations

import io
import re
import struct
import zlib
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

from . import LegacyConvertError

# ── 레코드 태그 (HWPTAG_BEGIN = 0x10) ───────────────────────────
_TAG_BEGIN = 0x10
TAG_ID_MAPPINGS = _TAG_BEGIN + 1       # 0x11 DocInfo
TAG_BIN_DATA = _TAG_BEGIN + 2          # 0x12 DocInfo
TAG_FACE_NAME = _TAG_BEGIN + 3         # 0x13 DocInfo
TAG_BORDER_FILL = _TAG_BEGIN + 4       # 0x14 DocInfo
TAG_CHAR_SHAPE = _TAG_BEGIN + 5        # 0x15 DocInfo
TAG_PARA_SHAPE = _TAG_BEGIN + 9        # 0x19 DocInfo
TAG_PARA_HEADER = _TAG_BEGIN + 50      # 0x42
TAG_PARA_TEXT = _TAG_BEGIN + 51        # 0x43
TAG_PARA_CHAR_SHAPE = _TAG_BEGIN + 52  # 0x44
TAG_CTRL_HEADER = _TAG_BEGIN + 55      # 0x47
TAG_LIST_HEADER = _TAG_BEGIN + 56      # 0x48
TAG_PAGE_DEF = _TAG_BEGIN + 57         # 0x49
TAG_SHAPE_COMPONENT = _TAG_BEGIN + 60  # 0x4C
TAG_TABLE = _TAG_BEGIN + 61            # 0x4D
TAG_SHAPE_PICTURE = _TAG_BEGIN + 69    # 0x55

#: 컨트롤 문자 크기 표 (pyhwp ControlChar 와 동일) — 코드 → 워드 수.
_CTRL_SIZES = {
    0x00: 1, 0x01: 8, 0x02: 8, 0x03: 8, 0x04: 8, 0x05: 8, 0x06: 8, 0x07: 8,
    0x08: 8, 0x09: 8, 0x0A: 1, 0x0B: 8, 0x0C: 8, 0x0D: 1, 0x0E: 8, 0x0F: 8,
    0x10: 8, 0x11: 8, 0x12: 8, 0x13: 8, 0x14: 8, 0x15: 8, 0x16: 8, 0x17: 8,
    0x18: 1, 0x1E: 1, 0x1F: 1,
}
_CTRL_RE = re.compile(rb"[\x00-\x1f]\x00")

_HWPUNIT_PER_INCH = 7200.0
_EMU_PER_INCH = 914400.0

#: python-docx 가 여는 이미지 컨테이너 — 그 외 확장자는 조용히 건너뛴다.
_DOCX_IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff"}


def _hu_to_emu(v: int) -> int:
    return int(round(v * _EMU_PER_INCH / _HWPUNIT_PER_INCH))


# ── DocInfo 카탈로그 ───────────────────────────────────────────


@dataclass
class _CharShape:
    """표 28 글자 모양 — 렌더에 실리는 부분집합."""
    size_pt: float = 10.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    color: Optional[str] = None  # RRGGBB
    face_id: Optional[int] = None  # 한글(ko) FaceName 참조


@dataclass
class _BorderSide:
    """표 20/21 테두리선 — stroke 0 = 선 없음."""
    stroke: int = 1
    width_mm: float = 0.12
    color: str = "000000"


@dataclass
class _BorderFill:
    """표 18 테두리/배경 — 좌/우/상/하 선 + 배경색."""
    bg: Optional[str] = None
    sides: Optional[List[_BorderSide]] = None  # [left, right, top, bottom]


@dataclass
class _ParaProps:
    """표 38 문단 모양 — 렌더에 실리는 부분집합."""
    align: str = "left"
    line_spacing: Optional[float] = None  # 배수 (RATIO 형만)
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0


@dataclass
class _DocInfo:
    char_shapes: List[_CharShape] = field(default_factory=list)
    para_shapes: List[_ParaProps] = field(default_factory=list)
    border_fills: List[_BorderFill] = field(default_factory=list)
    #: 한글(ko) 글꼴 이름 목록 — CHAR_SHAPE face_id 가 가리킨다.
    ko_faces: List[str] = field(default_factory=list)
    #: BIN_DATA 레코드 순서(1-based id) → (storage_id, ext). 링크형은 None.
    bin_data: List[Optional[Tuple[int, str]]] = field(default_factory=list)


def _read_bstr(payload: bytes, off: int) -> Tuple[str, int]:
    """BSTR: UINT16 글자 수 + UTF-16LE. (문자열, 다음 오프셋)."""
    if off + 2 > len(payload):
        return "", off
    (n,) = struct.unpack_from("<H", payload, off)
    end = off + 2 + n * 2
    if end > len(payload):
        return "", off
    return payload[off + 2:end].decode("utf-16le", errors="replace"), end


def _parse_char_shape(payload: bytes) -> _CharShape:
    """표 28/30 레이아웃: FontFace 7×WORD(14) + 폭/자간/상대크기/위치
    7×BYTE ×4(28) = 42, INT32 basesize(pt×100), UINT32 flags(bit0 italic,
    bit1 bold, bit2-3 밑줄 종류 — 1 밑줄/2 취소선/3 윗줄), COLORREF @52."""
    st = _CharShape()
    if len(payload) >= 2:
        (st.face_id,) = struct.unpack_from("<H", payload, 0)
    if len(payload) >= 46:
        (base,) = struct.unpack_from("<i", payload, 42)
        if 100 <= base <= 50000:
            st.size_pt = base / 100.0
    if len(payload) >= 50:
        (flags,) = struct.unpack_from("<I", payload, 46)
        st.italic = bool(flags & 0x1)
        st.bold = bool(flags & 0x2)
        kind = (flags >> 2) & 0x3
        st.underline = kind == 1
        st.strike = kind == 2
    if len(payload) >= 56:
        (colorref,) = struct.unpack_from("<I", payload, 52)
        r, g, b = colorref & 0xFF, (colorref >> 8) & 0xFF, (colorref >> 16) & 0xFF
        if (r, g, b) != (255, 255, 255):
            st.color = f"{r:02X}{g:02X}{b:02X}"
    return st


#: 표 39 문단 모양 속성1 align (bits 2-4) → docx 정렬.
_ALIGN_MAP = {0: "justify", 1: "left", 2: "right", 3: "center",
              4: "justify", 5: "justify"}


def _parse_para_shape(payload: bytes) -> _ParaProps:
    """표 38: flags(4) margins×4(doubled, 1/7200in ×2) @4..20,
    linespacing @24 (flags bits0-1: 0=RATIO %, 1=FIXED …)."""
    pp = _ParaProps()
    if len(payload) >= 4:
        (flags,) = struct.unpack_from("<I", payload, 0)
        pp.align = _ALIGN_MAP.get((flags >> 2) & 0x7, "left")
        if len(payload) >= 28:
            top2, bottom2, ls = struct.unpack_from("<3i", payload, 16)
            # doubled margin: 1/7200 inch × 2 → pt = v/2 × 72/7200 = v/200
            if 0 < top2 <= 7200 * 8:
                pp.space_before_pt = top2 / 200.0
            if 0 < bottom2 <= 7200 * 8:
                pp.space_after_pt = bottom2 / 200.0
            if (flags & 0x3) == 0 and 50 <= ls <= 500:  # RATIO(%)
                pp.line_spacing = ls / 100.0
    return pp


#: 표 21 테두리선 굵기 인덱스 → mm (pyhwp Border.widths)
_BORDER_WIDTH_MM = (0.1, 0.12, 0.15, 0.2, 0.25, 0.3, 0.4, 0.5,
                    0.6, 0.7, 1.0, 1.5, 2.0, 3.0, 4.0, 5.0)


def _parse_border_fill(payload: bytes) -> _BorderFill:
    """표 18: borderflags(2) + Border(stroke 1B, width 1B, COLORREF 4B)×5
    (좌/우/상/하/대각) = 32, fillflags UINT32 @32, colorpattern 이면
    background COLORREF @36."""
    bf = _BorderFill()
    if len(payload) >= 32:
        sides: List[_BorderSide] = []
        for k in range(4):  # left, right, top, bottom (대각선은 범위 밖)
            off = 2 + k * 6
            stroke = payload[off] & 0x1F
            width_idx = payload[off + 1] & 0x0F
            (colorref,) = struct.unpack_from("<I", payload, off + 2)
            r, g, b = colorref & 0xFF, (colorref >> 8) & 0xFF, (colorref >> 16) & 0xFF
            sides.append(_BorderSide(
                stroke=stroke,
                width_mm=_BORDER_WIDTH_MM[width_idx],
                color=f"{r:02X}{g:02X}{b:02X}"))
        bf.sides = sides
    if len(payload) >= 40:
        (fillflags,) = struct.unpack_from("<I", payload, 32)
        if fillflags & 0x1:
            (colorref,) = struct.unpack_from("<I", payload, 36)
            r, g, b = colorref & 0xFF, (colorref >> 8) & 0xFF, (colorref >> 16) & 0xFF
            if (r, g, b) != (255, 255, 255):
                bf.bg = f"{r:02X}{g:02X}{b:02X}"
    return bf


def _parse_bin_data(payload: bytes) -> Optional[Tuple[int, str]]:
    """표 12: flags UINT16 — EMBEDDING(1)/STORAGE(2)면 storage_id UINT16 +
    (EMBEDDING 은) ext BSTR. 링크형(0)은 외부 파일이라 None."""
    if len(payload) < 4:
        return None
    (flags,) = struct.unpack_from("<H", payload, 0)
    storage_type = flags & 0xF
    if storage_type not in (1, 2):
        return None
    (storage_id,) = struct.unpack_from("<H", payload, 2)
    ext = ""
    if storage_type == 1:
        ext, _ = _read_bstr(payload, 4)
    return storage_id, ext.lower().lstrip(".")


def _parse_doc_info(data: bytes) -> _DocInfo:
    info = _DocInfo()
    face_names: List[str] = []
    ko_face_count: Optional[int] = None
    for tagid, _level, payload in _iter_records(data):
        if tagid == TAG_ID_MAPPINGS:
            # 표 8: INT32 배열 — [0] binData, [1..7] 언어별 글꼴 수(ko 부터).
            n = len(payload) // 4
            if n >= 2:
                (ko_face_count,) = struct.unpack_from("<i", payload, 4)
        elif tagid == TAG_FACE_NAME:
            # 표 14: flags BYTE + name BSTR (뒤의 대체글꼴/패노즈는 불필요)
            name, _ = _read_bstr(payload, 1)
            face_names.append(name)
        elif tagid == TAG_BORDER_FILL:
            info.border_fills.append(_parse_border_fill(payload))
        elif tagid == TAG_CHAR_SHAPE:
            info.char_shapes.append(_parse_char_shape(payload))
        elif tagid == TAG_PARA_SHAPE:
            info.para_shapes.append(_parse_para_shape(payload))
        elif tagid == TAG_BIN_DATA:
            info.bin_data.append(_parse_bin_data(payload))
    # FACE_NAME 레코드는 언어 그룹 순서(ko 먼저)로 나온다 — ko 수만큼이
    # CHAR_SHAPE 의 ko face id 색인 공간이다. ID_MAPPINGS 가 없거나 수가
    # 어긋나면 전체 목록을 그대로 쓴다 (이름 참조라 틀려도 무해).
    if ko_face_count is not None and 0 < ko_face_count <= len(face_names):
        info.ko_faces = face_names[:ko_face_count]
    else:
        info.ko_faces = face_names
    return info


# ── 본문 모델 ──────────────────────────────────────────────────


@dataclass
class _Para:
    """레코드에서 복원한 문단."""
    text: str = ""
    #: (문자 위치, charshape id) — PARA_CHAR_SHAPE 그대로.
    shape_spans: List[tuple] = field(default_factory=list)
    parashape_id: Optional[int] = None
    #: 이 문단에 앵커된 블록(표/그림/글상자 문단들) — 등장 순서.
    attachments: List[object] = field(default_factory=list)


@dataclass
class _Cell:
    """표 75 셀 속성 + 내용 문단."""
    row: int = 0
    col: int = 0
    rowspan: int = 1
    colspan: int = 1
    width_hu: int = 0
    height_hu: int = 0
    borderfill_id: int = 0
    valign: str = "center"  # 표 60 listflags bits5-6 — 한글 기본은 가운데
    #: 표 75 안쪽 여백 (left,right,top,bottom) HWPUNIT16
    padding_hu: Tuple[int, int, int, int] = (0, 0, 0, 0)
    paras: List[_Para] = field(default_factory=list)


@dataclass
class _Table:
    rows: int = 0
    cols: int = 0
    cells: List[_Cell] = field(default_factory=list)
    caption: List[_Para] = field(default_factory=list)


@dataclass
class _Image:
    bindata_id: int = 0
    width_hu: int = 0
    height_hu: int = 0


@dataclass
class _TextBox:
    paras: List[_Para] = field(default_factory=list)


@dataclass
class _PageDef:
    width: int = 59528
    height: int = 84188
    left: int = 8504
    right: int = 8504
    top: int = 5668
    bottom: int = 4252


@dataclass
class _Node:
    tagid: int
    level: int
    payload: bytes
    children: List["_Node"] = field(default_factory=list)


# ── 저수준 파서 ────────────────────────────────────────────────


def _iter_records(data: bytes):
    """(tagid, level, payload) 를 순서대로 — 손상 시 그 지점에서 멈춘다."""
    pos, n = 0, len(data)
    while pos + 4 <= n:
        (hdr,) = struct.unpack_from("<I", data, pos)
        pos += 4
        tagid = hdr & 0x3FF
        level = (hdr >> 10) & 0x3FF
        size = (hdr >> 20) & 0xFFF
        if size == 0xFFF:
            if pos + 4 > n:
                return
            (size,) = struct.unpack_from("<I", data, pos)
            pos += 4
        if tagid == 0 and size == 0:
            return  # 패딩/끝
        if pos + size > n:
            return
        yield tagid, level, data[pos:pos + size]
        pos += size


def _build_tree(data: bytes) -> List[_Node]:
    """레코드 평면열 → level 기반 트리 (부모 = 직전의 더 얕은 레코드)."""
    roots: List[_Node] = []
    stack: List[_Node] = []
    for tagid, level, payload in _iter_records(data):
        node = _Node(tagid, level, payload)
        while stack and stack[-1].level >= level:
            stack.pop()
        (stack[-1].children if stack else roots).append(node)
        stack.append(node)
    return roots


def _decompress(raw: bytes) -> bytes:
    """HWP 압축 스트림 = raw deflate. 뒤에 패딩이 붙어 있어도 관용한다."""
    d = zlib.decompressobj(-15)
    out = d.decompress(raw)
    out += d.flush()
    return out


def _parse_text_chunks(payload: bytes) -> str:
    """PARA_TEXT → 본문 텍스트. 탭/줄바꿈은 보존, 컨트롤 워드는 건너뛴다."""
    parts: List[str] = []
    idx, n = 0, len(payload)
    while idx < n:
        m = _CTRL_RE.search(payload, idx)
        # 홀수 오프셋 매치는 UTF-16 상위바이트 우연 — 다음 짝수로.
        while m is not None and (m.start() & 1):
            m = _CTRL_RE.search(payload, m.start() + 1)
        ctrl = m.start() if m is not None else n
        if idx < ctrl:
            parts.append(payload[idx:ctrl].decode("utf-16le", errors="replace"))
        if m is None:
            break
        code = payload[ctrl]
        words = _CTRL_SIZES.get(code, 1)
        if code == 0x09:
            parts.append("\t")
        elif code == 0x0A:
            parts.append("\n")
        idx = ctrl + words * 2
    return "".join(parts)


def _parse_page_def(payload: bytes) -> _PageDef:
    pd = _PageDef()
    if len(payload) >= 24:
        w, h, left, right, top, bottom = struct.unpack_from("<6I", payload, 0)
        # 손상/이형 방어 — 종이 크기로 말이 되는 값만 채택 (1~50 inch).
        if 7200 <= w <= 360000 and 7200 <= h <= 360000:
            pd.width, pd.height = w, h
            pd.left, pd.right, pd.top, pd.bottom = left, right, top, bottom
    return pd


def _chid_of(payload: bytes) -> str:
    """CTRL_HEADER 의 chid — 리틀엔디언 4바이트 ('tbl ' 은 b' lbt')."""
    if len(payload) < 4:
        return ""
    return payload[:4][::-1].decode("ascii", errors="replace")


# ── 트리 해석 ──────────────────────────────────────────────────


def _parse_cell_props(payload: bytes) -> _Cell:
    """표 60 리스트 헤더(8B) 뒤에 표 75 셀 속성이 붙는다.

    셀 속성이 잘려 있으면(이형 파일) 좌표를 -1 로 표시한다 — 호출자가
    행 우선 순서로 재배정한다.
    """
    c = _Cell()
    if len(payload) >= 8:
        # 표 60 리스트 헤더 flags @4 — bits5-6 VAlign(0 top/1 middle/2 bottom)
        (listflags,) = struct.unpack_from("<I", payload, 4)
        c.valign = {0: "top", 1: "center", 2: "bottom"}.get(
            (listflags >> 5) & 0x3, "center")
    if len(payload) >= 16:
        c.col, c.row, c.colspan, c.rowspan = struct.unpack_from("<4H", payload, 8)
        c.colspan = max(1, c.colspan)
        c.rowspan = max(1, c.rowspan)
    else:
        c.col = c.row = -1
    if len(payload) >= 24:
        c.width_hu, c.height_hu = struct.unpack_from("<2i", payload, 16)
    if len(payload) >= 32:
        c.padding_hu = struct.unpack_from("<4H", payload, 24)
    if len(payload) >= 34:
        (c.borderfill_id,) = struct.unpack_from("<H", payload, 32)
    return c


def _list_paras(kids: List[_Node], i: int) -> Tuple[List[_Node], int]:
    """kids[i] (LIST_HEADER) 에 속한 문단 노드들과 다음 인덱스.

    **실파일에서 문단 리스트의 문단들은 LIST_HEADER 의 자식이 아니라 같은
    레벨의 형제다** (표 60 문단 리스트 헤더 — pyhwp table.hwp 실측:
    L2 LIST_HEADER 다음 L2 PARA_HEADER). 형제 수는 헤더의 paragraphs
    필드(UINT16 @0)가 상한이다. 자식으로 중첩된 이형(구버전 산출물)도
    함께 받아들인다.
    """
    lh = kids[i]
    para_nodes = [c for c in lh.children if c.tagid == TAG_PARA_HEADER]
    declared = 0
    if len(lh.payload) >= 2:
        (declared,) = struct.unpack_from("<H", lh.payload, 0)
    j = i + 1
    taken = 0
    while j < len(kids) and kids[j].tagid == TAG_PARA_HEADER:
        if declared > 0 and taken >= declared:
            break
        para_nodes.append(kids[j])
        taken += 1
        j += 1
    return para_nodes, j


def _interpret_table(ctrl: _Node) -> Optional[_Table]:
    """CTRL_HEADER('tbl ') 서브트리 → _Table.

    TABLE(표 70) 레코드가 격자 크기를, 그 **뒤의** LIST_HEADER 들이 셀을
    준다 (표 75 — col/row/colspan/rowspan 포함). TABLE **앞의** LIST_HEADER
    는 캡션이다 (한컴 스펙의 before/after tablebody 구분). 셀 내용 문단은
    LIST_HEADER 의 형제 — _list_paras 참조.
    """
    table = _Table()
    seen_body = False
    kids = ctrl.children
    i = 0
    while i < len(kids):
        child = kids[i]
        if child.tagid == TAG_TABLE:
            if len(child.payload) >= 8:
                rows, cols = struct.unpack_from("<HH", child.payload, 4)
                if 0 < rows <= 2000 and 0 < cols <= 256:
                    table.rows, table.cols = rows, cols
            seen_body = True
            i += 1
            continue
        if child.tagid == TAG_LIST_HEADER:
            para_nodes, i = _list_paras(kids, i)
            paras = [p for p in _interpret_paras(para_nodes)
                     if isinstance(p, _Para)]
            if seen_body:
                cell = _parse_cell_props(child.payload)
                cell.paras = paras
                # 셀 안의 표/그림도 셀 문단의 attachment 로 이미 들어 있다.
                table.cells.append(cell)
            else:
                table.caption.extend(paras)
            continue
        i += 1
    if not table.rows or not table.cols:
        return None
    # 셀 속성이 잘린 이형 파일 — 행 우선 순서로 좌표를 재배정한다.
    if any(c.col < 0 or c.row < 0 for c in table.cells):
        for i, c in enumerate(table.cells):
            c.row, c.col = i // table.cols, i % table.cols
            c.rowspan = c.colspan = 1
    return table


def _interpret_gso(ctrl: _Node) -> List[object]:
    """CTRL_HEADER('gso ') 서브트리 → [_Image | _TextBox] (등장 순서).

    SHAPE_COMPONENT(표 78)가 개체 크기를, SHAPE_COMPONENT_PICTURE(표 102)의
    PictureInfo.bindata_id(오프셋 71)가 이미지 원본을 가리킨다. 글상자는
    SHAPE_COMPONENT 아래 LIST_HEADER 의 문단들이다. 컨테이너는 재귀.
    """
    out: List[object] = []

    def walk(node: _Node, width_hu: int, height_hu: int) -> None:
        kids = node.children
        i = 0
        while i < len(kids):
            child = kids[i]
            if child.tagid == TAG_SHAPE_COMPONENT:
                w, h = width_hu, height_hu
                if len(child.payload) >= 36:
                    w2, h2 = struct.unpack_from("<2i", child.payload, 28)
                    if 0 < w2 <= 7200 * 100 and 0 < h2 <= 7200 * 100:
                        w, h = w2, h2
                walk(child, w, h)
            elif child.tagid == TAG_SHAPE_PICTURE:
                if len(child.payload) >= 73:
                    (bindata_id,) = struct.unpack_from("<H", child.payload, 71)
                    if bindata_id:
                        out.append(_Image(bindata_id, width_hu, height_hu))
            elif child.tagid == TAG_LIST_HEADER:
                para_nodes, i = _list_paras(kids, i)
                paras = [p for p in _interpret_paras(para_nodes)
                         if isinstance(p, _Para)]
                if paras:
                    out.append(_TextBox(paras))
                continue
            i += 1

    walk(ctrl, 0, 0)
    return out


def _interpret_paras(nodes: List[_Node]) -> List[object]:
    """PARA_HEADER 노드 목록 → [_Para] — 표/그림/글상자는 앵커 문단의
    attachments 로 붙는다 (본문 흐름상 그 문단 위치에서 등장).

    규격상 컨트롤은 앵커 문단의 자식이지만, 최상위에 직접 놓인 CTRL_HEADER
    (이형/편집기 산출물)도 빈 앵커 문단으로 감싸 받아들인다.
    """
    out: List[object] = []
    for node in nodes:
        if node.tagid == TAG_CTRL_HEADER:
            chid = _chid_of(node.payload)
            holder = _Para()
            if chid == "tbl ":
                table = _interpret_table(node)
                if table is not None:
                    holder.attachments.append(table)
            elif chid == "gso ":
                holder.attachments.extend(_interpret_gso(node))
            if holder.attachments:
                out.append(holder)
            continue
        if node.tagid != TAG_PARA_HEADER:
            continue
        para = _Para()
        if len(node.payload) >= 10:
            (para.parashape_id,) = struct.unpack_from("<H", node.payload, 8)
        for child in node.children:
            if child.tagid == TAG_PARA_TEXT:
                para.text += _parse_text_chunks(child.payload)
            elif child.tagid == TAG_PARA_CHAR_SHAPE:
                for off in range(0, len(child.payload) - 7, 8):
                    pos, shape_id = struct.unpack_from("<II", child.payload, off)
                    para.shape_spans.append((pos, shape_id))
            elif child.tagid == TAG_CTRL_HEADER:
                chid = _chid_of(child.payload)
                if chid == "tbl ":
                    table = _interpret_table(child)
                    if table is not None:
                        para.attachments.append(table)
                elif chid == "gso ":
                    para.attachments.extend(_interpret_gso(child))
        out.append(para)
    return out


def _find_records(nodes: List[_Node], tagid: int):
    """트리 전체에서 tagid 레코드를 깊이 우선으로 찾는다."""
    for node in nodes:
        if node.tagid == tagid:
            yield node
        yield from _find_records(node.children, tagid)


def _header_footer_paras(roots: List[_Node]) -> Tuple[List[_Para], List[_Para]]:
    """첫 머리말('head')/꼬리말('foot') 정의의 문단들."""
    header: List[_Para] = []
    footer: List[_Para] = []
    for ctrl in _find_records(roots, TAG_CTRL_HEADER):
        chid = _chid_of(ctrl.payload)
        if chid not in ("head", "foot"):
            continue
        target = header if chid == "head" else footer
        if target:
            continue  # 첫 정의만
        kids = ctrl.children
        i = 0
        while i < len(kids):
            if kids[i].tagid == TAG_LIST_HEADER:
                para_nodes, i = _list_paras(kids, i)
                target.extend(p for p in _interpret_paras(para_nodes)
                              if isinstance(p, _Para))
                continue
            i += 1
    return header, footer


# ── DOCX 조립 ──────────────────────────────────────────────────


def hwp_to_docx(content: bytes) -> bytes:
    import olefile
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt, RGBColor

    if not olefile.isOleFile(io.BytesIO(content)):
        raise LegacyConvertError("hwp 가 아닙니다 (OLE 복합문서 아님)")
    ole = olefile.OleFileIO(io.BytesIO(content))
    try:
        if not ole.exists("FileHeader"):
            raise LegacyConvertError("hwp FileHeader 스트림이 없습니다")
        header = ole.openstream("FileHeader").read()
        if not header.startswith(b"HWP Document File"):
            raise LegacyConvertError("hwp 시그니처가 아닙니다")
        (flags,) = struct.unpack_from("<I", header, 36)
        compressed = bool(flags & 0x1)
        if flags & 0x2:
            raise LegacyConvertError("암호로 보호된 hwp 는 열 수 없습니다")
        if flags & 0x4:
            raise LegacyConvertError("배포용(암호화) hwp 는 열 수 없습니다")

        def read_stream(name: str) -> bytes:
            raw = ole.openstream(name).read()
            return _decompress(raw) if compressed else raw

        info = _parse_doc_info(read_stream("DocInfo")) if ole.exists("DocInfo") \
            else _DocInfo()

        def bin_blob(bindata_id: int) -> Optional[Tuple[bytes, str]]:
            """PictureInfo.bindata_id(1-based) → (원본 바이트, 확장자)."""
            if not (1 <= bindata_id <= len(info.bin_data)):
                return None
            entry = info.bin_data[bindata_id - 1]
            if entry is None:
                return None
            storage_id, ext = entry
            if ext not in _DOCX_IMAGE_EXTS:
                return None
            name = f"BinData/BIN{storage_id:04X}.{ext}"
            try:
                if not ole.exists(name):
                    return None
                return read_stream(name), ext
            except Exception:  # noqa: BLE001 — 깨진 이미지는 건너뛴다
                return None

        # BodyText/Section* — 숫자 순
        section_names = sorted(
            ("/".join(entry) for entry in ole.listdir()
             if len(entry) == 2 and entry[0] == "BodyText"
             and entry[1].startswith("Section")),
            key=lambda s: int(re.sub(r"\D", "", s) or 0),
        )
        if not section_names:
            raise LegacyConvertError("hwp 본문(BodyText/Section*)이 없습니다")

        doc = Document()
        page_applied = False

        _WD_ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        def shape_of(idx: int) -> Optional[_CharShape]:
            return info.char_shapes[idx] if 0 <= idx < len(info.char_shapes) \
                else None

        def face_name(face_id: Optional[int]) -> Optional[str]:
            if face_id is None or not (0 <= face_id < len(info.ko_faces)):
                return None
            return info.ko_faces[face_id] or None

        def apply_align(para_obj, para: _Para) -> None:
            pid = para.parashape_id
            if pid is None or not (0 <= pid < len(info.para_shapes)):
                return
            pp = info.para_shapes[pid]
            if pp.align != "left":
                para_obj.alignment = _WD_ALIGN[pp.align]
            pf = para_obj.paragraph_format
            if pp.line_spacing is not None and abs(pp.line_spacing - 1.0) > 0.01:
                pf.line_spacing = pp.line_spacing
            if pp.space_before_pt > 0.05:
                pf.space_before = Pt(pp.space_before_pt)
            if pp.space_after_pt > 0.05:
                pf.space_after = Pt(pp.space_after_pt)

        def emit_runs(para_obj, para: _Para) -> None:
            text = para.text
            if not text:
                return
            spans = sorted(para.shape_spans) or [(0, -1)]
            for i, (start, shape_id) in enumerate(spans):
                end = spans[i + 1][0] if i + 1 < len(spans) else len(text)
                chunk = text[start:end]
                if not chunk:
                    continue
                st = shape_of(shape_id)
                # 줄바꿈은 run break 로 — 문단은 유지된다.
                pieces = chunk.split("\n")
                for j, piece in enumerate(pieces):
                    if j > 0:
                        para_obj.add_run().add_break()
                    if not piece:
                        continue
                    run = para_obj.add_run(piece)
                    if st is not None:
                        run.font.size = Pt(st.size_pt)
                        if st.bold:
                            run.bold = True
                        if st.italic:
                            run.italic = True
                        if st.underline:
                            run.underline = True
                        if st.strike:
                            run.font.strike = True
                        if st.color:
                            run.font.color.rgb = RGBColor.from_string(st.color)
                        name = face_name(st.face_id)
                        if name:
                            run.font.name = name
                            # 한글 글리프는 eastAsia 슬롯을 본다.
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn("w:eastAsia"), name)

        def set_cell_bg(cell_obj, rgb: str) -> None:
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), rgb)

        #: 표 20 stroke → docx 테두리 val (0 = 선 없음 → nil)
        _STROKE_VAL = {0: "nil", 1: "single", 2: "dashed", 3: "dotted",
                       4: "dotDash", 5: "dotDotDash", 6: "dashed",
                       7: "dotted", 8: "double", 9: "double", 10: "double",
                       11: "triple", 12: "wave", 13: "doubleWave"}

        def set_cell_borders(cell_obj, sides: List[_BorderSide]) -> None:
            """표 18 좌/우/상/하 → w:tcBorders (sz = 1/8pt)."""
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for name, side in zip(("left", "right", "top", "bottom"), sides):
                el = borders.find(qn(f"w:{name}"))
                if el is None:
                    el = OxmlElement(f"w:{name}")
                    borders.append(el)
                el.set(qn("w:val"), _STROKE_VAL.get(side.stroke, "single"))
                # mm → pt(×72/25.4) → 1/8pt
                el.set(qn("w:sz"), str(max(2, int(round(side.width_mm * 72 / 25.4 * 8)))))
                el.set(qn("w:color"), side.color)

        def set_cell_margins(cell_obj, padding_hu: Tuple[int, int, int, int]) -> None:
            """표 75 안쪽 여백 → w:tcMar (twips = hwpunit/5)."""
            if not any(padding_hu):
                return
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for name, hu in zip(("left", "right", "top", "bottom"), padding_hu):
                el = OxmlElement(f"w:{name}")
                el.set(qn("w:w"), str(int(hu / 5)))
                el.set(qn("w:type"), "dxa")
                mar.append(el)

        def set_cell_valign(cell_obj, valign: str) -> None:
            if valign == "top":
                return  # docx 기본
            tc_pr = cell_obj._tc.get_or_add_tcPr()
            va = tc_pr.find(qn("w:vAlign"))
            if va is None:
                va = OxmlElement("w:vAlign")
                tc_pr.append(va)
            va.set(qn("w:val"), valign)

        def fill_cell_paras(cell_obj, paras: List[_Para]) -> None:
            first = True
            for para in paras:
                if first:
                    cell_obj.paragraphs[0].text = ""
                    target = cell_obj.paragraphs[0]
                    first = False
                else:
                    target = cell_obj.add_paragraph()
                apply_align(target, para)
                emit_runs(target, para)
                emit_attachments(para, container_cell=cell_obj)

        def emit_table(block: _Table, container_cell=None) -> None:
            r_n, c_n = block.rows, block.cols
            if container_cell is not None:
                # 중첩 표 — 셀 안에 실제 표로 넣는다 (python-docx add_table).
                tbl = container_cell.add_table(rows=r_n, cols=c_n)
            else:
                tbl = doc.add_table(rows=r_n, cols=c_n)
            try:
                tbl.style = "Table Grid"
            except Exception:  # noqa: BLE001 — 스타일 없는 템플릿 관용
                pass

            # 열 너비: colspan 1 셀의 표 75 width 로 gridCol 을 채운다.
            col_w: Dict[int, int] = {}
            for cell in block.cells:
                if cell.colspan == 1 and 0 <= cell.col < c_n and cell.width_hu > 0:
                    col_w.setdefault(cell.col, cell.width_hu)
            for j, column in enumerate(tbl.columns):
                if j in col_w:
                    try:
                        column.width = Emu(_hu_to_emu(col_w[j]))
                    except Exception:  # noqa: BLE001
                        pass
            # 행 높이: rowspan 1 셀 height 의 최대값 (atLeast 의미).
            row_h: Dict[int, int] = {}
            for cell in block.cells:
                if cell.rowspan == 1 and 0 <= cell.row < r_n and cell.height_hu > 0:
                    row_h[cell.row] = max(row_h.get(cell.row, 0), cell.height_hu)
            for i, row in enumerate(tbl.rows):
                if i in row_h:
                    try:
                        row.height = Emu(_hu_to_emu(row_h[i]))
                    except Exception:  # noqa: BLE001
                        pass

            # 병합 먼저 (표 75 의 col/row/colspan/rowspan 은 절대 격자 좌표).
            for cell in block.cells:
                if cell.rowspan > 1 or cell.colspan > 1:
                    r2 = min(cell.row + cell.rowspan - 1, r_n - 1)
                    c2 = min(cell.col + cell.colspan - 1, c_n - 1)
                    if (r2, c2) != (cell.row, cell.col):
                        try:
                            tbl.cell(cell.row, cell.col).merge(tbl.cell(r2, c2))
                        except Exception:  # noqa: BLE001 — 겹침/범위 이상 관용
                            pass
            # 내용/배경.
            for cell in block.cells:
                if not (0 <= cell.row < r_n and 0 <= cell.col < c_n):
                    continue
                try:
                    cell_obj = tbl.cell(cell.row, cell.col)
                except Exception:  # noqa: BLE001
                    continue
                bf = None
                if 0 <= cell.borderfill_id - 1 < len(info.border_fills):
                    bf = info.border_fills[cell.borderfill_id - 1]
                if bf is not None:
                    if bf.bg:
                        set_cell_bg(cell_obj, bf.bg)
                    if bf.sides is not None:
                        set_cell_borders(cell_obj, bf.sides)
                set_cell_valign(cell_obj, cell.valign)
                set_cell_margins(cell_obj, cell.padding_hu)
                fill_cell_paras(cell_obj, cell.paras)
            for cap in block.caption:
                p = doc.add_paragraph() if container_cell is None \
                    else container_cell.add_paragraph()
                apply_align(p, cap)
                emit_runs(p, cap)

        def emit_image(img: _Image, container_cell=None) -> None:
            got = bin_blob(img.bindata_id)
            if got is None:
                return
            blob, _ext = got
            target = doc.add_paragraph() if container_cell is None \
                else container_cell.add_paragraph()
            run = target.add_run()
            try:
                kwargs = {}
                if img.width_hu > 0:
                    kwargs["width"] = Emu(_hu_to_emu(img.width_hu))
                if img.height_hu > 0:
                    kwargs["height"] = Emu(_hu_to_emu(img.height_hu))
                run.add_picture(io.BytesIO(blob), **kwargs)
            except Exception:  # noqa: BLE001 — 못 여는 이미지는 건너뛴다
                pass

        def emit_attachments(para: _Para, container_cell=None) -> None:
            for att in para.attachments:
                if isinstance(att, _Table):
                    emit_table(att, container_cell=container_cell)
                elif isinstance(att, _Image):
                    emit_image(att, container_cell=container_cell)
                elif isinstance(att, _TextBox):
                    for tb_para in att.paras:
                        p = doc.add_paragraph() if container_cell is None \
                            else container_cell.add_paragraph()
                        apply_align(p, tb_para)
                        emit_runs(p, tb_para)

        header_done = False
        for sec_name in section_names:
            roots = _build_tree(read_stream(sec_name))

            if not page_applied:
                for pd_node in _find_records(roots, TAG_PAGE_DEF):
                    page = _parse_page_def(pd_node.payload)
                    sec = doc.sections[0]
                    sec.page_width = Emu(_hu_to_emu(page.width))
                    sec.page_height = Emu(_hu_to_emu(page.height))
                    sec.left_margin = Emu(_hu_to_emu(page.left))
                    sec.right_margin = Emu(_hu_to_emu(page.right))
                    sec.top_margin = Emu(_hu_to_emu(page.top))
                    sec.bottom_margin = Emu(_hu_to_emu(page.bottom))
                    page_applied = True
                    break

            if not header_done:
                h_paras, f_paras = _header_footer_paras(roots)
                h_text = " ".join(p.text.strip() for p in h_paras if p.text.strip())
                f_text = " ".join(p.text.strip() for p in f_paras if p.text.strip())
                try:
                    if h_text:
                        doc.sections[0].header.paragraphs[0].text = h_text
                    if f_text:
                        doc.sections[0].footer.paragraphs[0].text = f_text
                except Exception:  # noqa: BLE001
                    pass
                header_done = bool(h_text or f_text)

            for block in _interpret_paras(roots):
                if not isinstance(block, _Para):
                    continue
                if block.text or not block.attachments:
                    p = doc.add_paragraph()
                    apply_align(p, block)
                    emit_runs(p, block)
                emit_attachments(block)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    finally:
        ole.close()
