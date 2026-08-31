"""HWPX(OWPML) → DOCX 구조 보존 변환.

HWPX 는 zip 컨테이너다:
    Contents/header.xml     charPr(글자 모양)·paraPr(문단 모양) 정의
    Contents/section*.xml   본문 — hp:p → hp:run → hp:t, 표(hp:tbl),
                            그림(hp:pic), 페이지 기하(hp:secPr/hp:pagePr)
    Contents/content.hpf    매니페스트 — binaryItemIDRef → BinData/* 경로
    BinData/*               이미지 원본

네임스페이스 프리픽스는 생산자마다 달라(hp/hwpml/…​) 로컬네임 매칭으로
파싱한다. 단위: HWPUNIT = 1/7200 inch, charPr height = 1/100 pt.
"""

from __future__ import annotations

import io
import posixpath
import re
import zipfile
from dataclasses import dataclass, field
from typing import Dict, List, Optional
from xml.etree import ElementTree as ET

from . import LegacyConvertError

_HWPUNIT_PER_INCH = 7200.0
_EMU_PER_INCH = 914400.0


def _hwpunit_to_emu(v: float) -> int:
    return int(round(v * _EMU_PER_INCH / _HWPUNIT_PER_INCH))


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _iter_local(elem: ET.Element, name: str):
    for child in elem.iter():
        if _local(child.tag) == name:
            yield child


def _find_local(elem: ET.Element, name: str) -> Optional[ET.Element]:
    for child in elem.iter():
        if _local(child.tag) == name:
            return child
    return None


def _children_local(elem: ET.Element, name: str):
    for child in list(elem):
        if _local(child.tag) == name:
            yield child


@dataclass
class _CharStyle:
    size_pt: float = 10.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: Optional[str] = None  # RRGGBB


@dataclass
class _ParaStyle:
    align: Optional[str] = None  # LEFT/CENTER/RIGHT/JUSTIFY


@dataclass
class _Page:
    width_hu: float = 59528.0   # A4 portrait
    height_hu: float = 84188.0
    margins_hu: Dict[str, float] = field(default_factory=lambda: {
        "left": 8504, "right": 8504, "top": 5668, "bottom": 4252,
    })


def _parse_header(header_xml: bytes) -> tuple[Dict[str, _CharStyle], Dict[str, _ParaStyle]]:
    chars: Dict[str, _CharStyle] = {}
    paras: Dict[str, _ParaStyle] = {}
    try:
        root = ET.fromstring(header_xml)
    except ET.ParseError as exc:
        raise LegacyConvertError(f"hwpx header.xml 파싱 실패: {exc}") from exc

    for pr in _iter_local(root, "charPr"):
        cid = pr.get("id")
        if cid is None:
            continue
        st = _CharStyle()
        height = pr.get("height")
        if height:
            try:
                st.size_pt = max(1.0, float(height) / 100.0)
            except ValueError:
                pass
        color = (pr.get("textColor") or "").lstrip("#")
        if re.fullmatch(r"[0-9A-Fa-f]{6}", color) and color.upper() != "FFFFFF":
            st.color = color.upper()
        for child in pr:
            name = _local(child.tag)
            if name == "bold":
                st.bold = True
            elif name == "italic":
                st.italic = True
            elif name == "underline" and (child.get("type") or "").upper() not in ("", "NONE"):
                st.underline = True
        chars[cid] = st

    for pr in _iter_local(root, "paraPr"):
        pid = pr.get("id")
        if pid is None:
            continue
        st = _ParaStyle()
        align_el = _find_local(pr, "align")
        if align_el is not None:
            st.align = (align_el.get("horizontal") or "").upper() or None
        paras[pid] = st
    return chars, paras


def _parse_manifest(zf: zipfile.ZipFile) -> Dict[str, str]:
    """binaryItemIDRef → zip 내 경로 (content.hpf 의 opf:item)."""
    out: Dict[str, str] = {}
    try:
        root = ET.fromstring(zf.read("Contents/content.hpf"))
    except (KeyError, ET.ParseError):
        return out
    for item in _iter_local(root, "item"):
        iid, href = item.get("id"), item.get("href")
        if iid and href:
            out[iid] = href
    return out


def _read_bindata(zf: zipfile.ZipFile, manifest: Dict[str, str], ref: str) -> Optional[bytes]:
    href = manifest.get(ref)
    candidates = [href] if href else []
    # 매니페스트가 없거나 어긋난 생산자 대비 — BinData/ 밑에서 이름 매칭.
    candidates += [n for n in zf.namelist()
                   if n.startswith("BinData/") and Path_stem(n) == ref]
    for cand in candidates:
        if not cand:
            continue
        for name in (cand, posixpath.join("Contents", cand)):
            try:
                return zf.read(name)
            except KeyError:
                continue
    return None


def Path_stem(name: str) -> str:
    base = name.rsplit("/", 1)[-1]
    return base.rsplit(".", 1)[0]


def _page_from_section(sec_root: ET.Element) -> _Page:
    page = _Page()
    page_pr = _find_local(sec_root, "pagePr")
    if page_pr is not None:
        try:
            page.width_hu = float(page_pr.get("width") or page.width_hu)
            page.height_hu = float(page_pr.get("height") or page.height_hu)
        except ValueError:
            pass
        margin = _find_local(page_pr, "margin")
        if margin is not None:
            for key in ("left", "right", "top", "bottom"):
                v = margin.get(key)
                if v:
                    try:
                        page.margins_hu[key] = float(v)
                    except ValueError:
                        pass
    return page


def hwpx_to_docx(content: bytes) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Pt, RGBColor

    try:
        zf = zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile as exc:
        raise LegacyConvertError("hwpx 가 아닙니다 (zip 컨테이너 아님)") from exc

    names = set(zf.namelist())
    section_names = sorted(n for n in names
                           if re.fullmatch(r"Contents/section\d+\.xml", n))
    if not section_names:
        raise LegacyConvertError("hwpx 본문(section*.xml)이 없습니다")

    char_styles, para_styles = (
        _parse_header(zf.read("Contents/header.xml"))
        if "Contents/header.xml" in names else ({}, {})
    )
    manifest = _parse_manifest(zf)

    doc = Document()
    # 기본 템플릿의 빈 문단 제거 방지용 — python-docx 새 문서는 빈 상태다.
    section = doc.sections[0]

    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "BOTH": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "DISTRIBUTE": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def apply_run_style(run, char_ref: Optional[str]) -> None:
        st = char_styles.get(char_ref or "")
        if st is None:
            return
        run.font.size = Pt(st.size_pt)
        if st.bold:
            run.bold = True
        if st.italic:
            run.italic = True
        if st.underline:
            run.underline = True
        if st.color:
            run.font.color.rgb = RGBColor.from_string(st.color)

    def emit_paragraph(p_el: ET.Element, container) -> None:
        """hp:p 하나 → docx 문단 (+ 내부의 표/그림은 뒤이어 블록으로)."""
        para = container.add_paragraph()
        st = para_styles.get(p_el.get("paraPrIDRef") or "")
        if st and st.align and st.align in align_map:
            para.alignment = align_map[st.align]

        pending_tables: List[ET.Element] = []
        pending_pics: List[ET.Element] = []
        for run_el in _children_local(p_el, "run"):
            char_ref = run_el.get("charPrIDRef")
            for child in list(run_el):
                name = _local(child.tag)
                if name == "t":
                    # hp:t 내부에 lineBreak 등의 자식이 섞일 수 있다.
                    if child.text:
                        apply_run_style(para.add_run(child.text), char_ref)
                    for sub in list(child):
                        if _local(sub.tag) == "lineBreak":
                            para.add_run().add_break()
                        if sub.tail:
                            apply_run_style(para.add_run(sub.tail), char_ref)
                elif name == "tbl":
                    pending_tables.append(child)
                elif name == "pic":
                    pending_pics.append(child)

        for tbl_el in pending_tables:
            emit_table(tbl_el, container)
        for pic_el in pending_pics:
            emit_picture(pic_el, container)

    def emit_table(tbl_el: ET.Element, container) -> None:
        rows = int(tbl_el.get("rowCnt") or 0)
        cols = int(tbl_el.get("colCnt") or 0)
        cells = []
        for tr in _children_local(tbl_el, "tr"):
            for tc in _children_local(tr, "tc"):
                cells.append(tc)
        if rows <= 0 or cols <= 0:
            return
        table = container.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        occupied = [[False] * cols for _ in range(rows)]
        for tc in cells:
            addr = _find_local(tc, "cellAddr")
            span = _find_local(tc, "cellSpan")
            r = int(addr.get("rowAddr") or 0) if addr is not None else 0
            c = int(addr.get("colAddr") or 0) if addr is not None else 0
            rs = int(span.get("rowSpan") or 1) if span is not None else 1
            cs = int(span.get("colSpan") or 1) if span is not None else 1
            if r >= rows or c >= cols:
                continue
            # 주소 미기재 생산자 폴백 — 다음 빈 칸 (행 우선).
            if addr is None:
                placed = False
                for rr in range(rows):
                    for cc in range(cols):
                        if not occupied[rr][cc]:
                            r, c = rr, cc
                            placed = True
                            break
                    if placed:
                        break
            cell = table.cell(r, c)
            if rs > 1 or cs > 1:
                r2 = min(rows - 1, r + rs - 1)
                c2 = min(cols - 1, c + cs - 1)
                cell = cell.merge(table.cell(r2, c2))
                for rr in range(r, r2 + 1):
                    for cc in range(c, c2 + 1):
                        occupied[rr][cc] = True
            else:
                occupied[r][c] = True
            sub = _find_local(tc, "subList")
            if sub is not None:
                first = True
                for p_el in _children_local(sub, "p"):
                    text_runs = [t.text or "" for t in _iter_local(p_el, "t")]
                    text = "".join(text_runs)
                    if first:
                        cell.paragraphs[0].text = ""
                        target = cell.paragraphs[0]
                        first = False
                    else:
                        target = cell.add_paragraph()
                    char_ref = None
                    for run_el in _children_local(p_el, "run"):
                        char_ref = run_el.get("charPrIDRef")
                        break
                    apply_run_style(target.add_run(text), char_ref)

    def emit_picture(pic_el: ET.Element, container) -> None:
        img_el = _find_local(pic_el, "img")
        ref = img_el.get("binaryItemIDRef") if img_el is not None else None
        if not ref:
            return
        data = _read_bindata(zf, manifest, ref)
        if not data:
            return
        width_emu = None
        sz = _find_local(pic_el, "sz")
        if sz is not None and sz.get("width"):
            try:
                width_emu = _hwpunit_to_emu(float(sz.get("width")))
            except ValueError:
                width_emu = None
        para = container.add_paragraph()
        run = para.add_run()
        try:
            run.add_picture(io.BytesIO(data),
                            width=Emu(width_emu) if width_emu else None)
        except Exception:  # noqa: BLE001 — 미지원 이미지 코덱은 건너뛴다
            return

    first_section_applied = False
    for sec_name in section_names:
        try:
            sec_root = ET.fromstring(zf.read(sec_name))
        except ET.ParseError as exc:
            raise LegacyConvertError(f"{sec_name} 파싱 실패: {exc}") from exc
        if not first_section_applied:
            page = _page_from_section(sec_root)
            from docx.shared import Emu as _Emu

            section.page_width = _Emu(_hwpunit_to_emu(page.width_hu))
            section.page_height = _Emu(_hwpunit_to_emu(page.height_hu))
            section.left_margin = _Emu(_hwpunit_to_emu(page.margins_hu["left"]))
            section.right_margin = _Emu(_hwpunit_to_emu(page.margins_hu["right"]))
            section.top_margin = _Emu(_hwpunit_to_emu(page.margins_hu["top"]))
            section.bottom_margin = _Emu(_hwpunit_to_emu(page.margins_hu["bottom"]))
            first_section_applied = True
        # 부모 맵 — 표 셀 내부 문단(subList 하위)은 표 렌더에서 처리하므로
        # 최상위 순회에서 제외한다 (ElementTree 에는 부모 포인터가 없다).
        parent = {c: par for par in sec_root.iter() for c in par}

        def _inside_sublist(el: ET.Element) -> bool:
            cur = parent.get(el)
            while cur is not None:
                if _local(cur.tag) == "subList":
                    return True
                cur = parent.get(cur)
            return False

        for p_el in _iter_local(sec_root, "p"):
            if _inside_sublist(p_el):
                continue
            emit_paragraph(p_el, doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
