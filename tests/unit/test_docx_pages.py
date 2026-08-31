"""M3 — DOCX page engine (docx_pages.docx_to_page_svgs)."""

from __future__ import annotations

from xgen_edit2docs.documents.docx_engine import docx_from_markdown
from xgen_edit2docs.documents.docx_pages import docx_to_page_svgs


def _doc(md: str) -> bytes:
    return docx_from_markdown(md)


class TestPagination:
    def test_single_page_for_short_doc(self):
        pages = docx_to_page_svgs(_doc("# Title\n\nOne paragraph."))
        assert len(pages) == 1
        assert pages[0].startswith("<svg")

    def test_long_doc_flows_to_multiple_pages(self):
        md = "\n\n".join(f"Paragraph {i}. " + "내용 텍스트 " * 20 for i in range(80))
        pages = docx_to_page_svgs(_doc(md))
        assert len(pages) >= 2

    def test_honors_sectpr_page_size(self):
        # python-docx's default template declares US Letter (12240×15840
        # twips → 816×1056 px) — the engine must read sectPr, not assume A4.
        page = docx_to_page_svgs(_doc("x"))[0]
        assert 'width="816"' in page and 'height="1056"' in page


class TestContent:
    def test_paragraph_addressing_tags(self):
        pages = docx_to_page_svgs(_doc("first\n\nsecond\n\nthird"))
        joined = "".join(pages)
        assert 'data-e2d-para="0"' in joined
        assert 'data-e2d-para="2"' in joined

    def test_table_renders_with_addresses(self):
        md = "| a | b |\n|---|---|\n| 1 | 2 |"
        joined = "".join(docx_to_page_svgs(_doc(md)))
        assert 'data-e2d-table="0"' in joined
        assert 'data-e2d-cell="1,1"' in joined
        assert joined.count("<rect") >= 4  # page bg + cells

    def test_heading_is_larger_and_bold(self):
        joined = "".join(docx_to_page_svgs(_doc("# 큰제목\n\n본문")))
        assert 'font-weight="bold"' in joined
        assert 'font-size="26.67"' in joined  # 20pt → 26.67px

    def test_bullets_render(self):
        joined = "".join(docx_to_page_svgs(_doc("- one\n- two")))
        assert "•" in joined

    def test_text_is_escaped(self):
        joined = "".join(docx_to_page_svgs(_doc("a < b & c > d")))
        assert "&lt;" in joined and "&amp;" in joined
        assert "a < b" not in joined


def _docx_bytes(doc) -> bytes:
    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestTableFidelity:
    """격자 모델 — 병합 기하/셀 스타일/정렬/행높이 (python-docx 로 직접 조립)."""

    def test_vmerge_draws_one_spanning_rect(self):
        import re

        from docx import Document

        doc = Document()
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).merge(t.cell(1, 0))
        t.cell(0, 0).text = "세로병합"
        t.cell(0, 1).text = "위"
        t.cell(1, 1).text = "아래"
        joined = "".join(docx_to_page_svgs(_docx_bytes(doc)))
        # 병합 앵커 텍스트는 한 번만, 그 셀의 rect 는 오른쪽 셀 두 개 높이의 합
        assert joined.count("세로병합") == 1
        heights = [float(h) for h in re.findall(r'<rect [^>]*height="([\d.]+)"', joined)]
        assert len(heights) >= 4  # 페이지 배경 + 셀 3개
        cell_hs = sorted(heights[1:])
        assert cell_hs[-1] >= cell_hs[0] * 1.9  # 앵커 rect ≈ 2행 높이

    def test_cell_runs_keep_styles_and_alignment(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        p = t.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("빨강 굵게")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FF0000")
        run2 = p.add_run("취소선")
        run2.font.strike = True
        joined = "".join(docx_to_page_svgs(_docx_bytes(doc)))
        assert 'font-weight="bold"' in joined
        assert "#FF0000" in joined
        assert "line-through" in joined

    def test_body_paragraph_alignment(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        left = doc.add_paragraph("왼쪽정렬문장")
        center = doc.add_paragraph("가운데정렬문장")
        center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        joined = "".join(docx_to_page_svgs(_docx_bytes(doc)))

        def x_of(text: str) -> float:
            import re

            m = re.search(rf'<text x="([\d.]+)"[^>]*>{text}', joined)
            assert m, text
            return float(m.group(1))

        assert x_of("가운데정렬문장") > x_of("왼쪽정렬문장") + 50

    def test_trheight_minimum_is_honored(self):
        import re

        from docx import Document
        from docx.shared import Twips

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "x"
        t.rows[0].height = Twips(3000)  # 3000/15 = 200px
        joined = "".join(docx_to_page_svgs(_docx_bytes(doc)))
        heights = [float(h) for h in re.findall(r'<rect [^>]*height="([\d.]+)"', joined)]
        assert any(abs(h - 200.0) < 1.0 for h in heights)

    def test_nested_table_content_not_lost(self):
        from docx import Document

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        inner = t.cell(0, 0).add_table(rows=1, cols=2)
        inner.cell(0, 0).text = "안쪽A"
        inner.cell(0, 1).text = "안쪽B"
        joined = "".join(docx_to_page_svgs(_docx_bytes(doc)))
        assert "안쪽A" in joined and "안쪽B" in joined
